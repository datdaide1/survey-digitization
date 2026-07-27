"""Dựng báo cáo DOCX gửi khách — đủ Tầng 1-7: tổng quan, Phần A/B1/B2/B3 theo danh mục
biểu đồ §11 docs/implement-plan-statistics-and-client-report.md, văn phong §9.

Ảnh biểu đồ là PNG tĩnh (scripts/lib/report/mpl_charts.py) — khác "Biểu đồ" sheet
trong XLSX (native Excel chart, sống). Số liệu trong bảng/câu văn tính bằng
scripts/lib/report/frequency.py (Python/pandas) — độc lập với công thức Excel, dùng để
đối chiếu (spot-check) 2 nguồn cho khớp trước khi gửi khách.
"""

from __future__ import annotations

import tempfile
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from . import mpl_charts as charts
from . import narrative
from .association_sheet import SECTION_TITLES
from .cluster_analysis import ClusterResult
from .codebook import build_codebook
from .curated_pairs import resolve_label
from .factor_analysis import FactorAnalysisResult
from .frequency import descriptive_stats, frequency_table
from .reliability import Q14_ROWS, Q32_ROWS

CODEBOOK = build_codebook()


def _add_freq_table(doc: Document, freq: dict[str, Any]) -> None:
    rows = [r for r in freq["rows"] if r["n"] > 0] or freq["rows"]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "Lựa chọn", "Số phiếu", "Tỷ lệ %"
    for r in rows:
        cells = table.add_row().cells
        cells[0].text = str(r["label"])
        cells[1].text = str(r["n"])
        cells[2].text = f"{r['pct']:.1f}%"
    doc.add_paragraph(f"n hợp lệ = {freq['valid_n']}, thiếu/không rõ = {freq['missing_n']}").runs[0].italic = True


def _add_picture(doc: Document, png_path: str, width_cm: float = 15) -> None:
    doc.add_picture(png_path, width=Cm(width_cm))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def _pie_or_bar(freq: dict[str, Any], out_path) -> str:
    """Pie giả định các lát cộng lại đúng 100% số phiếu — sai với câu có `has_overlap`
    (1 phiếu tick ≥2 ô trên câu chỉ cho chọn 1, vd Q12/Q13/Q20/Q27a — xem frequency.py).
    Tự chuyển sang bar ngang (cách biểu diễn đúng cho dữ liệu có thể chồng lấn), giữ
    nguyên pie cho câu thực sự chỉ 1 đáp án/phiếu."""
    if freq.get("has_overlap"):
        return charts.bar_chart(freq, out_path)
    return charts.pie_chart(freq, out_path)


def _multi_select_rows(df: pd.DataFrame, qid: str, options: list[str]) -> tuple[list[dict], int]:
    # 26/07: mẫu số = TỔNG SỐ PHIẾU (len(df), cố định 85), không phải số phiếu đã trả
    # lời câu này (trước đây dùng notna() của cột đầu tiên — bị "phồng" % lên khi có
    # phiếu bỏ trống cả câu). Cùng nguyên tắc với frequency_table(), xem frequency.py.
    cols = [f"{qid}_{o}" for o in options]
    n_total = len(df)
    rows = []
    for opt, col in zip(options, cols):
        label = CODEBOOK[col]["label"].split(" — ", 1)[-1]
        n = int((df[col] == 1).sum())
        pct = (n / n_total * 100) if n_total else 0.0
        rows.append({"label": label, "n": n, "pct": pct})
    return rows, n_total


def _matrix_row_summary(df: pd.DataFrame, qid: str, row_codes: list[str], col_labels_by_code: dict[str, str]) -> tuple[list[str], list[str], list[dict]]:
    # 26/07: mẫu số mỗi dòng = TỔNG SỐ PHIẾU (len(df), cố định 85), không phải số phiếu
    # có trả lời dòng đó (trước đây dùng n_valid = dropna() riêng từng dòng). Phiếu bỏ
    # trống 1 dòng cụ thể giờ vẫn tính vào mẫu số 85 — % các cột trong 1 dòng có thể
    # cộng lại < 100% nếu dòng đó có phiếu bỏ trống (không phải lỗi).
    #
    # 26/07 (phản hồi khách): trước đây 1 phiếu tick ≥2 ô trên 1 dòng ma trận (ghi dạng
    # combo 'vo+chong') bị dồn hết vào cột "Khác/nhiều lựa chọn" thay vì cộng vào TỪNG cột
    # nó thực sự tick — sai nguyên tắc so với cách multi-select thật/frequency_table() xử
    # lý combo (xem frequency.py._combo_parts): 1 phiếu tick cả vợ & chồng phải cộng dồn
    # vào CẢ hàng 'Vợ' lẫn hàng 'Chồng', không phải tạo thành 1 hạng mục riêng. Cột
    # "Khác/nhiều lựa chọn" giờ CHỈ còn chứa giá trị lạ không khớp code nào đã biết (nếu
    # có) — không còn chứa combo của các code đã biết nữa.
    row_labels = []
    pct_matrix = []
    col_labels = list(col_labels_by_code.values()) + ["Khác/nhiều lựa chọn"]
    total_n = len(df)
    known_codes = set(col_labels_by_code)
    for row_code in row_codes:
        col = f"{qid}_{row_code}"
        row_labels.append(CODEBOOK[col]["label"].split(" — ", 1)[-1])
        valid = df[col].dropna()
        counts: dict[str, int] = {code: 0 for code in col_labels_by_code}
        other_n = 0
        for value in valid:
            parts = value.split("+") if isinstance(value, str) and "+" in value else None
            if parts and all(p in known_codes for p in parts):
                for p in parts:
                    counts[p] += 1
            elif value in known_codes:
                counts[value] += 1
            else:
                other_n += 1
        pct_row: dict[str, float] = {
            label: (counts[code] / total_n * 100 if total_n else 0.0)
            for code, label in col_labels_by_code.items()
        }
        pct_row["Khác/nhiều lựa chọn"] = (other_n / total_n * 100) if total_n else 0.0
        pct_matrix.append(pct_row)

    # 26/07 (phản hồi khách): sau khi combo đã được cộng dồn đúng vào từng lựa chọn thật ở
    # trên, cột "Khác/nhiều lựa chọn" hầu như luôn = 0% (Q14/Q32 hiện tại: luôn = 0, không
    # còn giá trị lạ nào trong dữ liệu) — vẫn hiện trên biểu đồ/chú giải dù trống trơn gây
    # khó hiểu ("sao vẫn còn cột này?"). Chỉ giữ cột này nếu THẬT SỰ có giá trị > 0 ở ít
    # nhất 1 dòng; bỏ hẳn khỏi kết quả nếu không, tự động thích ứng nếu sau này dữ liệu mới
    # phát sinh giá trị lạ.
    if all(row["Khác/nhiều lựa chọn"] == 0 for row in pct_matrix):
        col_labels = col_labels[:-1]
        for row in pct_matrix:
            del row["Khác/nhiều lựa chọn"]

    return row_labels, col_labels, pct_matrix


def build_docx(
    df: pd.DataFrame,
    out_path: str | Path,
    tmp_dir: str | None = None,
    association_info: dict | None = None,
    effect_size_info: dict | None = None,
    reliability_info: dict | None = None,
    advanced_info: dict | None = None,
) -> None:
    doc = Document()
    tmp = Path(tmp_dir) if tmp_dir else Path(tempfile.mkdtemp(prefix="docx_charts_"))
    tmp.mkdir(parents=True, exist_ok=True)

    doc.add_heading("BÁO CÁO KHẢO SÁT SINH KẾ PHỤ NỮ TRỒNG CÂY DƯỢC LIỆU", level=0)
    doc.add_paragraph(
        "Báo cáo tổng hợp kết quả khảo sát 85 phụ nữ dân tộc thiểu số vùng cao trồng cây "
        "dược liệu tại hai tỉnh Lào Cai và Lai Châu. Số liệu trong báo cáo đã được ẩn danh, "
        "không còn tên hay thông tin liên hệ của người trả lời."
    )
    doc.add_paragraph(
        "Lưu ý: một phần dữ liệu vẫn đang trong quá trình rà soát lần cuối trước khi bàn giao "
        "chính thức — các con số trong bản này có thể được cập nhật lại sau khi rà soát xong."
    ).runs[0].italic = True

    doc.add_heading("Tổng quan nhanh", level=1)
    overview = doc.add_table(rows=1, cols=2)
    overview.style = "Light List Accent 1"
    overview.rows[0].cells[0].text, overview.rows[0].cells[1].text = "Chỉ số", "Giá trị"
    prov = frequency_table("province", df["province"])
    for r in prov["rows"]:
        cells = overview.add_row().cells
        prov_label = "Lào Cai" if r["code"] == "lao-cai" else "Lai Châu"
        cells[0].text, cells[1].text = f"Số phiếu — {prov_label}", f"{r['n']}"
    age_desc = descriptive_stats("Q2_tuoi", df["Q2_tuoi"])
    overview.add_row().cells[0].text = "Tuổi trung bình"
    overview.rows[-1].cells[1].text = f"{age_desc['mean']:.0f} tuổi"

    # --- Phần A ---
    doc.add_heading("Phần A — Thông tin chung", level=1)

    doc.add_heading("Q2 – Tuổi", level=2)
    desc = descriptive_stats("Q2_tuoi", df["Q2_tuoi"])
    doc.add_paragraph(narrative.continuous_sentence(desc))
    _add_picture(doc, charts.histogram(desc, df["Q2_tuoi"], tmp / "q2_tuoi.png"))
    age_bracket_freq = frequency_table("age_bracket", df["age_bracket"])
    doc.add_paragraph(narrative.categorical_sentence(age_bracket_freq))
    _add_picture(doc, charts.bar_chart(age_bracket_freq, tmp / "age_bracket.png"))
    _add_freq_table(doc, age_bracket_freq)

    # 26/07: bỏ thống kê Q3 (giới tính) — 100% phiếu là nữ (đối tượng khảo sát của
    # nghiên cứu vốn chỉ nhắm tới phụ nữ), 1 biến hằng số không mang thông tin gì để báo
    # cáo, tiện thể tránh biến variance=0 làm nhiễu ma trận tương quan/factor analysis ở
    # các tầng sau (đã bỏ khỏi phạm vi thống kê ngay từ flatten_stats, xem flatten.py).
    doc.add_heading("Q4 – Dân tộc", level=2)
    freq = frequency_table("Q4", df["Q4"])
    doc.add_paragraph(narrative.categorical_sentence(freq))
    _add_picture(doc, charts.bar_chart(freq, tmp / "q4.png"))
    _add_freq_table(doc, freq)

    doc.add_heading("Q5 – Trình độ học vấn", level=2)
    doc.add_paragraph(
        "Cấp học tính theo LỚP ĐÃ HỌC XONG, không phải lớp đang/từng học dở: học xong "
        "lớp 9 mới tính là hết THCS (học xong lớp 6, 7 hay 8 vẫn tính là hết tiểu học); "
        "học xong lớp 12 mới tính là hết THPT (học xong lớp 10 hay 11 vẫn tính là hết "
        "THCS; có học tiếp trung cấp/cao đẳng/đại học cũng tính vào nhóm này vì phải "
        "tốt nghiệp lớp 12 mới học tiếp được). Người chưa từng đi học chính quy ngày nào "
        "tính riêng, không gộp vào tiểu học."
    )
    # 26/07: gộp cả 4 mức cấp học (education_grade_bracket, đã gộp sẵn 2 tick trung_cap_dh
    # -> thpt và khong_di_hoc -> bucket riêng, xem bucketing.py) + 1 tick còn lại
    # (Q5_khong_tieng_pho_thong, không suy ra được cấp học nên vẫn để độc lập) vào
    # CHUNG 1 biểu đồ duy nhất — không tách chart-cấp-học/câu-văn-riêng như trước.
    edu_freq = frequency_table("education_grade_bracket", df["education_grade_bracket"])
    edu_rows_by_code = {r["code"]: r for r in edu_freq["rows"]}
    tieng_freq = frequency_table("Q5_khong_tieng_pho_thong", df["Q5_khong_tieng_pho_thong"])
    tieng_true = next(r for r in tieng_freq["rows"] if r["code"] is True)
    combined_rows = [
        edu_rows_by_code["khong_di_hoc"],
        edu_rows_by_code["tieu_hoc"],
        edu_rows_by_code["thcs"],
        edu_rows_by_code["thpt"],
        {"label": "Không nói được tiếng phổ thông (tiếng Kinh)", "n": tieng_true["n"], "pct": tieng_true["pct"]},
    ]
    doc.add_paragraph(
        f"Trong tổng {edu_freq['total_n']} phiếu khảo sát: "
        + "; ".join(f"{r['label'].lower()} {r['pct']:.0f}% ({r['n']} người)" for r in combined_rows)
        + "."
    )
    combined_edu_freq = {
        "label": "Q5 – Trình độ học vấn",
        "kind": "categorical",
        "valid_n": edu_freq["total_n"],
        "missing_n": 0,
        "total_n": edu_freq["total_n"],
        "rows": combined_rows,
        "has_overlap": True,  # "không nói được tiếng phổ thông" độc lập, có thể trùng bất kỳ mức cấp học nào
    }
    _add_picture(doc, charts.bar_chart(combined_edu_freq, tmp / "education.png"))
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "Lựa chọn", "Số phiếu", "Tỷ lệ %"
    for r in combined_rows:
        cells = table.add_row().cells
        cells[0].text, cells[1].text, cells[2].text = str(r["label"]), str(r["n"]), f"{r['pct']:.1f}%"

    doc.add_heading("Q6 – Hôn nhân", level=2)
    q6_freq = frequency_table("Q6", df["Q6"])
    doc.add_paragraph(narrative.categorical_sentence(q6_freq))
    _add_picture(doc, _pie_or_bar(q6_freq, tmp / "q6.png"))
    marriage_desc = descriptive_stats("Q6_tuoi_ket_hon", df["Q6_tuoi_ket_hon"])
    if marriage_desc["n"]:
        _add_picture(doc, charts.histogram(marriage_desc, df["Q6_tuoi_ket_hon"], tmp / "marriage_age.png"))
    marriage_bracket_freq = frequency_table("marriage_age_bracket", df["marriage_age_bracket"])
    doc.add_paragraph(narrative.marriage_age_sentence(marriage_bracket_freq))
    _add_picture(doc, charts.bar_chart(marriage_bracket_freq, tmp / "marriage_bracket.png"))
    _add_freq_table(doc, marriage_bracket_freq)

    doc.add_heading("Q7 – Nguồn thu nhập chính", level=2)
    q7_options = ["trong_trot", "chan_nuoi", "cay_duoc_lieu", "lam_nghiep", "phi_nong_nghiep"]
    q7_rows, q7_n = _multi_select_rows(df, "Q7", q7_options)
    doc.add_paragraph(narrative.multi_select_sentence("Q7 – Nguồn thu nhập chính", q7_rows, q7_n))
    _add_picture(doc, charts.multi_select_bar(
        "Q7 – Nguồn thu nhập chính", [r["label"] for r in q7_rows], [r["pct"] for r in q7_rows], q7_n, tmp / "q7.png"
    ))

    doc.add_heading("Q8 – Tỷ lệ thu nhập từ cây dược liệu", level=2)
    q8_freq = frequency_table("Q8", df["Q8"])
    doc.add_paragraph(narrative.categorical_sentence(q8_freq))
    _add_picture(doc, charts.bar_chart(q8_freq, tmp / "q8.png"))
    _add_freq_table(doc, q8_freq)

    doc.add_heading("Q9 – Số năm kinh nghiệm trồng dược liệu", level=2)
    exp_desc = descriptive_stats("Q9_derived_years_exp", df["Q9_derived_years_exp"])
    exp_bracket_freq = frequency_table("experience_years_bracket", df["experience_years_bracket"])
    doc.add_paragraph(narrative.experience_sentence(exp_desc, exp_bracket_freq))
    _add_picture(doc, charts.histogram(exp_desc, df["Q9_derived_years_exp"], tmp / "q9_exp.png"))
    _add_picture(doc, charts.bar_chart(exp_bracket_freq, tmp / "exp_bracket.png"))
    _add_freq_table(doc, exp_bracket_freq)

    doc.add_heading("Q10 – Nghề nghiệp chính", level=2)
    q10_freq = frequency_table("Q10", df["Q10"])
    doc.add_paragraph(narrative.categorical_sentence(q10_freq))
    _add_picture(doc, charts.bar_chart(q10_freq, tmp / "q10.png"))
    _add_freq_table(doc, q10_freq)

    doc.add_heading("Q11 – Hội đoàn thể", level=2)
    q11_options = ["hoi_phu_nu", "doan_thanh_nien", "hoi_nong_dan", "khong_hoi_vien"]
    q11_rows, q11_n = _multi_select_rows(df, "Q11", q11_options)
    doc.add_paragraph(narrative.multi_select_sentence("Q11 – Hội đoàn thể", q11_rows, q11_n))
    _add_picture(doc, charts.multi_select_bar(
        "Q11 – Hội đoàn thể", [r["label"] for r in q11_rows], [r["pct"] for r in q11_rows], q11_n, tmp / "q11.png"
    ))

    # --- Phần B1 — Phân công lao động ---
    doc.add_heading("Phần B1 — Phân công lao động", level=1)
    for qid, title in [("Q12", "Ai làm chính việc nhà"), ("Q13", "Ai tham gia chính trồng/bán dược liệu")]:
        doc.add_heading(f"{qid} – {title}", level=2)
        freq = frequency_table(qid, df[qid])
        doc.add_paragraph(narrative.categorical_sentence(freq))
        _add_picture(doc, _pie_or_bar(freq, tmp / f"{qid}.png"))

    doc.add_heading("Q14 – Phân công 18 công việc trong gia đình", level=2)
    q14_rows = [
        "lam_dat", "trong", "cham_soc_cay", "thuoc_bvtv", "thu_hoach", "so_che",
        "lien_he_tieu_thu", "khac_san_xuat", "noi_tro", "giat_giu", "dua_don_con",
        "dam_cuoi_gio", "cham_soc_con", "gia_suc_nho", "gia_suc_lon",
        "quan_ly_chi_tieu", "day_do_con", "bao_duong_xe",
    ]
    q14_col_labels = {"vo": "Vợ", "chong": "Chồng", "ca_hai": "Cả hai", "con_cai_lon": "Con cái lớn", "nguoi_khac": "Người khác"}
    row_labels, col_labels, pct_matrix = _matrix_row_summary(df, "Q14", q14_rows, q14_col_labels)
    doc.add_paragraph(
        "Biểu đồ dưới đây cho thấy trong mỗi việc, phần trăm phiếu ghi nhận việc đó do vợ, "
        "chồng, cả hai, con cái lớn, hay người khác trong nhà đảm nhận chính."
    )
    _add_picture(doc, charts.stacked_matrix_bar("Q14 – Ai làm chính từng việc trong gia đình", row_labels, col_labels, pct_matrix, tmp / "q14.png"), width_cm=17)

    doc.add_heading("Q16a – Thay đổi vai trò giới", level=2)
    q16a_freq = frequency_table("Q16a", df["Q16a"])
    doc.add_paragraph(narrative.categorical_sentence(q16a_freq))
    _add_picture(doc, _pie_or_bar(q16a_freq, tmp / "q16a.png"))

    # --- Phần B2 — Tiếp cận nguồn lực ---
    doc.add_heading("Phần B2 — Tiếp cận nguồn lực", level=1)

    doc.add_heading("Q17 – Sở hữu thiết bị", level=2)
    # 26/07: 2 tầng thống kê tách riêng theo yêu cầu khách — (1) tỷ lệ SỞ HỮU từng thiết
    # bị, mẫu số = toàn bộ phiếu trả lời câu này; (2) CHỈ trong nhóm có sở hữu thiết bị đó,
    # tỷ lệ do chồng/vợ đứng tên (mẫu số đổi thành n hộ sở hữu — có thể cả hai cùng tick nên
    # 2 % có thể cộng lại > 100%). Khác bản cũ: bản cũ tính %chồng/%vợ trên TOÀN BỘ mẫu,
    # không lọc theo sở hữu trước — sai với yêu cầu "có sở hữu mới tính ai sở hữu".
    q17_rows_codes = ["dien_thoai", "may_tinh", "may_tinh_bang"]
    q17_row_labels_map = {"dien_thoai": "Điện thoại thông minh", "may_tinh": "Máy tính", "may_tinh_bang": "Máy tính bảng"}
    q17_series = ["chong", "vo"]
    q17_series_labels_map = {"chong": "Chồng", "vo": "Vợ"}

    # 26/07: tầng (1) tỷ lệ sở hữu đổi mẫu số về TỔNG 85 PHIẾU (cố định), không phải chỉ
    # số phiếu có trả lời câu này (trước dùng valid_mask.notna()) — cùng nguyên tắc
    # frequency_table(). Tầng (2) chồng/vợ trong nhóm sở hữu VẪN giữ mẫu số = n_owned
    # (chủ đích lọc theo sở hữu, đây là quyết định riêng ở trên, không liên quan gì đến
    # việc loại phiếu bỏ trống).
    total_n = len(df)
    q17_own_pct = []
    q17_values_matrix = []
    for row_code in q17_rows_codes:
        chong_col = df[f"Q17_{row_code}_chong"]
        vo_col = df[f"Q17_{row_code}_vo"]
        owned_mask = (chong_col == 1) | (vo_col == 1)
        n_owned = int(owned_mask.sum())
        q17_own_pct.append((n_owned / total_n * 100) if total_n else 0.0)

        row_pct = {}
        for s in q17_series:
            col = df[f"Q17_{row_code}_{s}"]
            n_yes_among_owners = int((col[owned_mask] == 1).sum())
            row_pct[q17_series_labels_map[s]] = (n_yes_among_owners / n_owned * 100) if n_owned else 0.0
        q17_values_matrix.append(row_pct)

    q17_n_total = total_n
    # 26/07 (phản hồi khách): "Không ai trong nhà có bất kỳ thiết bị nào" trước đây chỉ in
    # thành 1 câu văn riêng, không lên biểu đồ — giờ gộp thẳng vào biểu đồ tỷ lệ sở hữu
    # làm 1 cột nữa, cùng dữ liệu, khỏi phải đọc số lẻ tẻ ngoài biểu đồ.
    q17_khong_freq = frequency_table("Q17_khong_ai_co", df["Q17_khong_ai_co"])
    q17_khong_row = next(r for r in q17_khong_freq["rows"] if r["code"] is True)
    q17_own_labels = [q17_row_labels_map[c] for c in q17_rows_codes] + ["Không ai trong nhà có thiết bị nào"]
    q17_own_pct_full = q17_own_pct + [q17_khong_row["pct"]]
    doc.add_paragraph(
        "Biểu đồ đầu tiên: tỷ lệ hộ có sở hữu từng loại thiết bị, và tỷ lệ hộ không có bất "
        "kỳ thiết bị nào (trên tổng 85 phiếu khảo sát). Biểu đồ thứ hai: CHỈ tính trong số "
        "hộ đã xác nhận có sở hữu thiết bị đó, tỷ lệ do chồng đứng tên và do vợ đứng tên (2 "
        "tỷ lệ có thể cộng lại quá 100% vì một thiết bị có thể do cả hai cùng đứng tên)."
    )
    _add_picture(doc, charts.multi_select_bar(
        "Q17 – Tỷ lệ hộ sở hữu thiết bị", q17_own_labels, q17_own_pct_full, q17_n_total, tmp / "q17_own.png"
    ))
    _add_picture(doc, charts.grouped_bar(
        "Q17 – Trong nhóm có sở hữu: do chồng/vợ đứng tên",
        [q17_row_labels_map[c] for c in q17_rows_codes],
        [q17_series_labels_map[s] for s in q17_series],
        q17_values_matrix, tmp / "q17.png",
    ))

    doc.add_heading("Q18/Q19 – Mục đích dùng thiết bị / ứng dụng công nghệ", level=2)
    q18_options = ["goi_nhan_tin", "doc_tin_tuc", "giai_tri", "tim_kiem_thong_tin", "giao_dich_ban_hang", "lam_viec_hoc_online", "mang_xa_hoi", "khac"]
    q18_rows, q18_n = _multi_select_rows(df, "Q18", q18_options)
    doc.add_paragraph(narrative.multi_select_sentence("Q18 – Mục đích dùng thiết bị", q18_rows, q18_n))
    _add_picture(doc, charts.multi_select_bar("Q18 – Mục đích dùng thiết bị", [r["label"] for r in q18_rows], [r["pct"] for r in q18_rows], q18_n, tmp / "q18.png"))

    q19_options = ["quang_ba", "quan_ly_sx_tc", "ban_hang_online", "khong"]
    q19_rows, q19_n = _multi_select_rows(df, "Q19", q19_options)
    doc.add_paragraph(narrative.multi_select_sentence("Q19 – Ứng dụng công nghệ trong sản xuất/kinh doanh", q19_rows, q19_n))
    _add_picture(doc, charts.multi_select_bar("Q19 – Ứng dụng công nghệ trong sản xuất/kinh doanh", [r["label"] for r in q19_rows], [r["pct"] for r in q19_rows], q19_n, tmp / "q19.png"))

    # 26/07 (phản hồi khách): trước đây khối này lên báo cáo lộn xộn — Q20, Q21a, Q23,
    # Q26, Q27a nằm 1 vòng lặp riêng, Q22a chèn sau đó, Q24/Q25 chèn sau nữa (không theo
    # số thứ tự câu hỏi), ĐỒNG THỜI Q21b, Q22b, Q28, Q29a, Q29b có trong schema/dữ liệu
    # nhưng chưa từng được viết vào DOCX. Giờ đi đúng thứ tự Q20 -> Q29b, không bỏ sót câu
    # nào (xem docs/implement-plan-statistics-and-client-report.md không giới hạn phạm vi
    # DOCX chỉ bằng phạm vi 11 chart XLSX Phần A).
    doc.add_heading("Q20 – Ai tham gia hội họp, tập huấn", level=2)
    q20_freq = frequency_table("Q20", df["Q20"])
    doc.add_paragraph(narrative.categorical_sentence(q20_freq))
    _add_picture(doc, _pie_or_bar(q20_freq, tmp / "q20.png"))

    doc.add_heading("Q21a – Tần suất tham gia tập huấn (2 năm qua)", level=2)
    q21a_freq = frequency_table("Q21a", df["Q21a"])
    doc.add_paragraph(narrative.categorical_sentence(q21a_freq))
    _add_picture(doc, charts.bar_chart(q21a_freq, tmp / "q21a.png"))

    doc.add_heading("Q21b – Nội dung tập huấn", level=2)
    q21b_options = ["ky_thuat", "khac"]
    q21b_rows, q21b_n = _multi_select_rows(df, "Q21b", q21b_options)
    doc.add_paragraph(narrative.multi_select_sentence("Q21b – Nội dung tập huấn", q21b_rows, q21b_n))
    _add_picture(doc, charts.multi_select_bar("Q21b – Nội dung tập huấn", [r["label"] for r in q21b_rows], [r["pct"] for r in q21b_rows], q21b_n, tmp / "q21b.png"))

    doc.add_heading("Q22a – Vay vốn", level=2)
    q22a_options = ["ngan_hang_thuong_mai", "ngan_hang_chinh_sach", "hoi_doan_the", "khac", "chua"]
    q22a_rows, q22a_n = _multi_select_rows(df, "Q22a", q22a_options)
    doc.add_paragraph(narrative.multi_select_sentence("Q22a – Vay vốn 5 năm qua", q22a_rows, q22a_n))
    _add_picture(doc, charts.multi_select_bar("Q22a – Vay vốn 5 năm qua", [r["label"] for r in q22a_rows], [r["pct"] for r in q22a_rows], q22a_n, tmp / "q22a.png"))

    doc.add_heading("Q22b – Lý do chưa vay vốn", level=2)
    q22b_options = ["khong_co_nhu_cau", "thu_tuc_phuc_tap", "khong_tai_san_the_chap", "khac"]
    q22b_rows, q22b_n = _multi_select_rows(df, "Q22b", q22b_options)
    doc.add_paragraph(narrative.multi_select_sentence("Q22b – Lý do chưa vay vốn", q22b_rows, q22b_n))
    _add_picture(doc, charts.multi_select_bar("Q22b – Lý do chưa vay vốn", [r["label"] for r in q22b_rows], [r["pct"] for r in q22b_rows], q22b_n, tmp / "q22b.png"))

    doc.add_heading("Q23 – Nhận hỗ trợ vật chất (5 năm qua)", level=2)
    q23_freq = frequency_table("Q23", df["Q23"])
    doc.add_paragraph(narrative.categorical_sentence(q23_freq))
    _add_picture(doc, _pie_or_bar(q23_freq, tmp / "q23.png"))

    doc.add_heading("Q24/Q25 – Đi lại tự chủ (xe máy)", level=2)
    # 26/07: mẫu số đổi về TỔNG 85 PHIẾU (cố định) — trước dùng mobility_valid (chỉ đếm
    # phiếu trả lời ĐỦ CẢ Q24 và Q25). Đếm số phiếu khớp từng nhóm vẫn dùng
    # mobility_valid (an toàn, NaN không khớp code nào nên không cần đổi), chỉ mẫu số %
    # là đổi.
    mobility_valid = df[["Q24", "Q25"]].dropna()
    mobility_groups = [
        ("co", "co", "Biết đi xe máy + có xe riêng"),
        ("co", "khong", "Biết đi xe máy + không có xe riêng"),
        ("khong", "co", "Không biết đi xe máy + có xe riêng"),
        ("khong", "khong", "Không biết đi xe máy + không có xe riêng"),
    ]
    mobility_valid_n = len(mobility_valid)
    mobility_n = len(df)
    mobility_rows = [
        {
            "label": label,
            "n": int(((mobility_valid["Q24"] == q24_code) & (mobility_valid["Q25"] == q25_code)).sum()),
            "pct": 0.0,
        }
        for q24_code, q25_code, label in mobility_groups
    ]
    for r in mobility_rows:
        r["pct"] = (r["n"] / mobility_n * 100) if mobility_n else 0.0
    mobility_freq = {
        "label": "Q24/Q25 – 4 nhóm tự chủ đi lại (xe máy)",
        "valid_n": mobility_valid_n,
        "missing_n": mobility_n - mobility_valid_n,
        "total_n": mobility_n,
        "rows": mobility_rows,
    }
    doc.add_paragraph(narrative.categorical_sentence(mobility_freq))
    _add_picture(doc, charts.bar_chart(mobility_freq, tmp / "mobility.png"))

    doc.add_heading("Q26 – Phải xin phép khi tham gia hoạt động xã hội", level=2)
    q26_freq = frequency_table("Q26", df["Q26"])
    doc.add_paragraph(narrative.categorical_sentence(q26_freq))
    _add_picture(doc, _pie_or_bar(q26_freq, tmp / "q26.png"))

    doc.add_heading("Q27a – Ai đứng tên quyền sử dụng đất", level=2)
    q27a_freq = frequency_table("Q27a", df["Q27a"])
    doc.add_paragraph(narrative.categorical_sentence(q27a_freq))
    _add_picture(doc, charts.bar_chart(q27a_freq, tmp / "q27a.png"))

    doc.add_heading("Q28 – Khó khăn khi trồng/kinh doanh dược liệu", level=2)
    q28_options = [
        "thieu_nguon_luc", "thieu_kien_thuc", "ganh_nang_viec_nha", "khong_thi_truong",
        "tap_quan_dinh_kien", "khac", "khong_kho_khan",
    ]
    q28_rows, q28_n = _multi_select_rows(df, "Q28", q28_options)
    doc.add_paragraph(narrative.multi_select_sentence("Q28 – Khó khăn khi trồng/kinh doanh dược liệu", q28_rows, q28_n))
    _add_picture(doc, charts.multi_select_bar("Q28 – Khó khăn khi trồng/kinh doanh dược liệu", [r["label"] for r in q28_rows], [r["pct"] for r in q28_rows], q28_n, tmp / "q28.png"))

    doc.add_heading("Q29a – Lợi ích từ trồng/bán dược liệu", level=2)
    q29a_freq = frequency_table("Q29a", df["Q29a"])
    doc.add_paragraph(narrative.categorical_sentence(q29a_freq))
    _add_picture(doc, _pie_or_bar(q29a_freq, tmp / "q29a.png"))

    doc.add_heading("Q29b – Lợi ích cụ thể", level=2)
    q29b_options = ["them_thu_nhap", "nang_cao_kien_thuc", "tham_gia_cong_dong"]
    q29b_rows, q29b_n = _multi_select_rows(df, "Q29b", q29b_options)
    doc.add_paragraph(narrative.multi_select_sentence("Q29b – Lợi ích cụ thể", q29b_rows, q29b_n))
    _add_picture(doc, charts.multi_select_bar("Q29b – Lợi ích cụ thể", [r["label"] for r in q29b_rows], [r["pct"] for r in q29b_rows], q29b_n, tmp / "q29b.png"))

    # --- Phần B3 — Ra quyết định ---
    doc.add_heading("Phần B3 — Ra quyết định", level=1)

    doc.add_heading("Q30 – Khâu tham gia nhiều nhất trong chuỗi giá trị", level=2)
    q30_freq = frequency_table("Q30", df["Q30"])
    doc.add_paragraph(narrative.categorical_sentence(q30_freq))
    _add_picture(doc, charts.bar_chart(q30_freq, tmp / "q30.png"))

    doc.add_heading("Q32 – Ai quyết định 8 vấn đề trong gia đình", level=2)
    q32_rows = ["chon_cay_trong", "mua_vat_tu", "chon_ban", "vay_von", "gia_ban", "su_dung_thu_nhap", "su_dung_dat", "khac"]
    q32_col_labels = {"vo": "Vợ", "chong": "Chồng", "cung_quyet_dinh": "Cùng quyết định", "con_cai": "Con cái"}
    row_labels, col_labels, pct_matrix = _matrix_row_summary(df, "Q32", q32_rows, q32_col_labels)
    _add_picture(doc, charts.stacked_matrix_bar("Q32 – Ai quyết định từng vấn đề trong gia đình", row_labels, col_labels, pct_matrix, tmp / "q32.png"), width_cm=17)

    # 26/07: bỏ mục "So sánh 8 vấn đề giữa 2 tỉnh" (radar theo tỉnh) — quyết định của
    # user, không cần so sánh giữa các vùng/tỉnh nữa. Cùng đợt bỏ: xlsx Cross-tab sheet +
    # Q32 radar native Excel (build_client_report.py), GROUP_COMPARISON_PAIRS['province']
    # (curated_pairs.py), so sánh trung bình theo tỉnh trong Chỉ số tổng hợp
    # (reliability_sheet.py). Xem [[project-survey-no-region-comparison]].

    doc.add_heading("Q33 – Vai trò lãnh đạo", level=2)
    q33_options = ["ban_chu_nhiem_htx", "nhom_san_xuat", "quan_ly_rung", "khong"]
    q33_rows, q33_n = _multi_select_rows(df, "Q33", q33_options)
    doc.add_paragraph(narrative.multi_select_sentence("Q33 – Vai trò lãnh đạo", q33_rows, q33_n))
    _add_picture(doc, charts.multi_select_bar("Q33 – Vai trò lãnh đạo", [r["label"] for r in q33_rows], [r["pct"] for r in q33_rows], q33_n, tmp / "q33.png"))

    if association_info is not None:
        _add_association_section(doc, association_info, tmp)

    if effect_size_info is not None:
        _add_effect_size_section(doc, effect_size_info, tmp)

    if reliability_info is not None:
        _add_reliability_section(doc, reliability_info, df, tmp)

    if advanced_info is not None:
        _add_advanced_section(doc, advanced_info, tmp)

    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def _add_reliability_section(doc: Document, reliability_info: dict, df: pd.DataFrame, tmp: Path) -> None:
    doc.add_heading("Tầng 6 — Chỉ số tổng hợp", level=1)
    doc.add_paragraph(
        "Thay vì đọc từng dòng riêng lẻ của 2 câu ma trận (phân công lao động, ra quyết "
        "định), phần này kiểm tra xem các dòng đó có đủ nhất quán để gộp thành 1 chỉ số "
        "tổng hợp duy nhất hay không (Cronbach's alpha, mốc tham khảo 0,7)."
    )
    for qid, r in reliability_info.items():
        doc.add_heading(r.label, level=2)
        if r.alpha is None:
            doc.add_paragraph("Không đủ dữ liệu để tính.")
            continue
        if r.composite is not None:
            doc.add_paragraph(
                f"Các dòng trong mục này đủ nhất quán để gộp thành 1 chỉ số chung "
                f"(alpha = {r.alpha:.2f}). Tính trung bình trên {len(r.composite)} phiếu, "
                f"chỉ số này đạt {r.composite.mean():.0f}/100 — càng cao nghĩa là vợ càng "
                f"tham gia/quyết định nhiều trong các việc/vấn đề được hỏi."
            )
            desc = r.composite.describe()
            doc.add_paragraph(
                f"Thấp nhất {desc['min']:.0f}, cao nhất {desc['max']:.0f}, phổ biến nhất quanh "
                f"mức {desc['50%']:.0f}."
            )
            png = charts.composite_distribution(
                f"{r.label} — phân bố trên {len(r.composite)} phiếu", r.composite.values, tmp / f"composite_{qid}.png"
            )
            _add_picture(doc, png, width_cm=13)
        else:
            doc.add_paragraph(
                f"Các dòng trong mục này chưa đủ nhất quán để gộp thành 1 chỉ số chung "
                f"(alpha = {r.alpha:.2f}, dưới mốc tham khảo 0,7) — báo cáo vẫn giữ nguyên "
                f"từng dòng riêng như phần Phân công lao động/Ra quyết định ở trên."
            )


def _effect_sentence(row: dict) -> str:
    r = row["result"]
    a_label, b_label = resolve_label(row["a"], CODEBOOK), resolve_label(row["b"], CODEBOOK)
    if r.value is None:
        return f"{a_label} và {b_label}: không đủ dữ liệu để tính."
    if row["kind"] == "odds_ratio":
        direction = "cao hơn" if r.value > 1 else "thấp hơn"
        times = r.value if r.value > 1 else (1 / r.value if r.value else None)
        times_txt = f"khoảng {times:.1f} lần" if times else ""
        return f"{row['desc']}: quan sát được khả năng {direction} {times_txt} (n={r.n})."
    if row["kind"] == "eta_squared":
        return f"{row['desc']}: khác biệt giữa các nhóm giải thích được khoảng {r.value*100:.0f}% biến thiên quan sát được (n={r.n})."
    return f"{row['desc']}: liên hệ quan sát khoảng {r.value:+.3f} đơn vị mỗi 1 đơn vị tăng thêm (n={r.n})."


def _add_effect_size_section(doc: Document, effect_size_info: dict, tmp: Path) -> None:
    doc.add_heading("Tầng 4 — Độ ảnh hưởng quan sát được", level=1)
    doc.add_paragraph(
        "Các con số dưới đây mô tả liên hệ QUAN SÁT ĐƯỢC giữa 2 câu trả lời, có kèm chiều "
        "(bên nào cao/thấp hơn) — KHÔNG PHẢI bằng chứng nhân quả, vì đây là khảo sát hỏi 1 "
        "lần, không có nhóm đối chứng hay thử nghiệm. Không nên đọc là 'A làm B thay đổi'."
    )
    for row in effect_size_info["effect_rows"]:
        doc.add_paragraph(_effect_sentence(row), style="List Bullet")

    # 26/07 (phản hồi khách): Tầng 4 trước đây chỉ có chữ (bullet list), không có biểu đồ
    # nào — thêm 2 biểu đồ tóm tắt trực quan, nhóm theo loại effect size (odds ratio và
    # eta-squared có thang đo khác nhau, không gộp chung 1 biểu đồ được; "slope" mỗi cặp 1
    # đơn vị đo riêng nên vẫn để dạng câu văn, không hợp để vẽ chung).
    or_rows = [row for row in effect_size_info["effect_rows"] if row["kind"] == "odds_ratio" and row["result"].value is not None]
    if or_rows:
        or_labels = [f"{resolve_label(row['a'], CODEBOOK)} → {resolve_label(row['b'], CODEBOOK)}" for row in or_rows]
        or_values = [row["result"].value for row in or_rows]
        png = charts.odds_ratio_forest("Tầng 4 — Odds ratio (mốc trung lập = 1)", or_labels, or_values, tmp / "effect_or.png")
        _add_picture(doc, png, width_cm=16)

    eta_rows = [row for row in effect_size_info["effect_rows"] if row["kind"] == "eta_squared" and row["result"].value is not None]
    if eta_rows:
        eta_labels = [f"{resolve_label(row['a'], CODEBOOK)} ↔ {resolve_label(row['b'], CODEBOOK)}" for row in eta_rows]
        eta_values = [row["result"].value * 100 for row in eta_rows]
        png = charts.eta_squared_bar("Tầng 4 — % biến thiên giải thích được (eta-squared)", eta_labels, eta_values, tmp / "effect_eta.png")
        _add_picture(doc, png, width_cm=16)

    doc.add_heading("Tầng 5 — So sánh giữa các nhóm", level=1)
    doc.add_paragraph(
        "Phần này so sánh xem một số nhóm (theo việc đứng tên đất, theo dân "
        "tộc...) có khác nhau về một số chỉ tiêu hay không. Với nhóm có ít phiếu (dưới "
        "10 phiếu), kết quả chỉ mang tính tham khảo, không nên xem là kết luận chắc chắn."
    )
    # 26/07 (phản hồi khách: "cơ sở nào để chọn các nhóm này mà so, không có số liệu nào,
    # toàn tự chọn xong tự tính"): các cặp so sánh dưới đây KHÔNG phải kết quả rà soát tự
    # động toàn bộ dữ liệu để tìm khác biệt lớn nhất — đây là ví dụ minh hoạ kỹ thuật so
    # sánh nhóm phi tham số, chọn sẵn cùng khách lúc lên kế hoạch báo cáo (không phải quét
    # hết mọi tổ hợp biến số × biến nhóm có thể có, việc đó sẽ ra hàng trăm cặp không phải
    # cặp nào cũng có ý nghĩa thực tế để đọc). Nói rõ ở đây để không hiểu nhầm là danh sách
    # đầy đủ/khách quan hoàn toàn — nếu cần so sánh thêm nhóm nào khác, có thể yêu cầu bổ
    # sung cụ thể.
    doc.add_paragraph(
        "Lưu ý về cách chọn: 2 phép so sánh dưới đây là ví dụ minh hoạ được chọn khi lên kế "
        "hoạch báo cáo (không phải máy quét tự động tìm khác biệt lớn nhất trong toàn bộ dữ "
        "liệu) — nếu cần so sánh thêm theo nhóm khác (vd theo xã, theo nhóm tuổi...), có thể "
        "yêu cầu bổ sung cụ thể."
    ).runs[0].italic = True
    _GROUP_LABELS = {
        "lao-cai": "Lào Cai", "lai-chau": "Lai Châu",
        "True": "Có", "False": "Không", "true": "Có", "false": "Không",
    }
    for i, gr in enumerate(effect_size_info["group_results"]):
        r = gr["result"]
        doc.add_heading(gr["desc"], level=2)
        if r.p_value is None:
            doc.add_paragraph("Không đủ dữ liệu để so sánh.")
            continue
        parts = [f"{_GROUP_LABELS.get(str(g['group']), str(g['group']))} (n={g['n']}, trung vị={g['median']})" for g in r.groups]
        small_groups = [g for g in r.groups if g["n"] < 10]
        sentence = "So sánh giữa các nhóm: " + "; ".join(parts) + "."
        if r.p_value < 0.05:
            sentence += " Khác biệt giữa các nhóm khó xảy ra do ngẫu nhiên thuần tuý trong mẫu này."
        else:
            sentence += " Khác biệt giữa các nhóm chưa đủ rõ để phân biệt với biến động ngẫu nhiên trong mẫu."
        if small_groups:
            sentence += " Lưu ý một số nhóm có rất ít phiếu, kết quả chỉ mang tính tham khảo."
        doc.add_paragraph(sentence)

        values_by_group = [{"label": str(g["group"]), "values": g["values"]} for g in r.groups if g.get("values")]
        if values_by_group:
            png = charts.box_whisker(gr["desc"], values_by_group, tmp / f"box_{i}.png")
            _add_picture(doc, png, width_cm=13)


def _add_association_section(doc: Document, association_info: dict, tmp: Path) -> None:
    doc.add_heading("Tầng 3 — Liên hệ giữa các câu trả lời", level=1)
    doc.add_paragraph(
        "Phần này xem những câu trả lời nào thường đi cùng nhau trong cùng một phiếu — "
        "ví dụ nhóm tuổi nào thường có cách phân công công việc khác nhóm tuổi khác, hay "
        "gia đình có xe máy/điện thoại riêng có thường tham gia tập huấn nhiều hơn không. "
        "Đây chỉ là quan sát mô tả từ 85 phiếu khảo sát, KHÔNG phải kết luận có ý nghĩa "
        "thống kê chặt chẽ hay quan hệ nhân quả — không nên diễn giải là 'A gây ra B'. "
        "(Không đưa liên hệ tỉnh↔dân tộc vào đây vì đó là kết quả của cách chọn mẫu theo "
        "xã — mỗi xã vốn đã tập trung sẵn 1-2 dân tộc — không phải điều mới quan sát được.)"
    )

    section_columns = association_info["section_columns"]
    matrix = association_info["matrix"]
    block_names = list(section_columns.keys())
    block_labels = [SECTION_TITLES[b] for b in block_names]
    summary = []
    for b1 in block_names:
        row = {}
        for b2 in block_names:
            vals = []
            for a in section_columns[b1]:
                for b in section_columns[b2]:
                    if a == b:
                        continue
                    r = matrix.get((a, b)) or matrix.get((b, a))
                    if r is not None and r.value is not None:
                        vals.append(abs(r.value))
            row[b2] = (sum(vals) / len(vals)) if vals else 0.0
        summary.append(row)

    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    fig, ax = plt.subplots(figsize=(7, 5.5))
    grid = [[summary[i][b2] for b2 in block_names] for i in range(len(block_names))]
    im = ax.imshow(grid, cmap="Blues", vmin=0, vmax=1)
    ax.set_xticks(range(len(block_names)))
    ax.set_xticklabels(block_labels, rotation=20, ha="right")
    ax.set_yticks(range(len(block_names)))
    ax.set_yticklabels(block_labels)
    for i in range(len(block_names)):
        for j in range(len(block_names)):
            ax.text(j, i, f"{grid[i][j]:.2f}", ha="center", va="center", color="black")
    ax.set_title("Mức độ liên quan trung bình giữa các khối câu hỏi")
    fig.colorbar(im, ax=ax, shrink=0.8)
    fig.tight_layout()
    png_path = tmp / "association_summary.png"
    fig.savefig(png_path, bbox_inches="tight")
    plt.close(fig)
    _add_picture(doc, str(png_path), width_cm=14)
    doc.add_paragraph(
        f"Khối A: {SECTION_TITLES['A'].split('— ', 1)[-1]}. "
        f"Khối B1: {SECTION_TITLES['B1'].split('— ', 1)[-1]}. "
        f"Khối B2: {SECTION_TITLES['B2'].split('— ', 1)[-1]}. "
        f"Khối B3: {SECTION_TITLES['B3'].split('— ', 1)[-1]}."
    ).runs[0].italic = True

    doc.add_heading("Các cặp liên quan rõ nhất", level=2)
    # 26/07 (phản hồi khách — bảng cũ đọc không hiểu đang thể hiện gì): tách rõ từng thông
    # tin ra cột riêng thay vì dồn hết "chiều + n + p" vào 1 câu "Ghi chú". Đồng thời SỬA 1
    # lỗi diễn giải: Cramér's V/eta (dùng cho hầu hết các cặp — biến phân loại không có
    # thứ tự, vd tỉnh/nghề nghiệp/dân tộc) KHÔNG có khái niệm "chiều tăng/giảm" — trước đây
    # bảng cứ in "cùng tăng/cùng có" cho MỌI cặp kể cả loại này, gây hiểu nhầm là có chiều
    # trong khi con số đó chỉ là ĐỘ MẠNH (0 đến 1), không có dấu. Chỉ Spearman/rank-biserial
    # (biến số hoặc biến 2 mức) mới thật sự có dấu +/- để nói "chiều".
    doc.add_paragraph(
        "Mỗi dòng là 1 cặp câu trả lời hay xuất hiện cùng nhau trong cùng 1 phiếu. 'Độ mạnh "
        "liên hệ' từ 0 (không liên hệ) đến 1 (liên hệ hoàn toàn) — cách tính (Cramér's V/"
        "Spearman/rank-biserial/eta) tuỳ loại biến, xem ghi chú đầu Tầng 3. 'Chiều' chỉ có ở "
        "cặp biến số/biến 2 mức (vd tuổi, có/không); cặp biến phân loại nhiều mức không thứ "
        "tự (vd tỉnh, dân tộc, nghề nghiệp) không có khái niệm 'chiều tăng/giảm' nên để "
        "trống. 'Đáng tin?' dựa trên p-value (mốc 0,05) — bảng NÀY đã loại sẵn các cặp có "
        "nhóm quá nhỏ (<5 phiếu) nên không còn dòng 'Chưa rõ'. Vẫn KHÔNG phải bằng chứng "
        "nhân quả — chỉ nói lên độ tin cậy của liên hệ quan sát được."
    )
    table = doc.add_table(rows=1, cols=6)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for c, text in zip(hdr, ["Câu trả lời 1", "Câu trả lời 2", "Độ mạnh liên hệ", "Chiều", "n", "Đáng tin? (p-value)"]):
        c.text = text
    SIGNED_METHODS = {"spearman", "rank_biserial"}
    for a, b, r in association_info["top20"][:15]:
        cells = table.add_row().cells
        cells[0].text = CODEBOOK[a]["label"]
        cells[1].text = CODEBOOK[b]["label"]
        cells[2].text = f"{abs(r.value):.2f}" if r.value is not None else "—"
        if r.method in SIGNED_METHODS:
            cells[3].text = "Cùng tăng/cùng có" if (r.value or 0) > 0 else "Ngược chiều"
        else:
            cells[3].text = "—"
        cells[4].text = str(r.n)
        sig = r.significant
        p_txt = f"p={r.p_value:.3f}" if r.p_value is not None else "p=?"
        cells[5].text = f"Có, {p_txt}" if sig is True else (f"Không rõ rệt, {p_txt}" if sig is False else "Chưa rõ")


def _item_label(qid: str, row_code: str) -> str:
    return CODEBOOK[f"{qid}_{row_code}"]["label"].split(" — ", 1)[-1]


def _factor_groupings(result: FactorAnalysisResult, rows: list[str], threshold: float = 0.4) -> list[list[str]]:
    """Với mỗi nhân tố, lấy các item có |loading| lớn nhất trên đúng nhân tố đó (tránh 1
    item xuất hiện ở nhiều nhóm) VÀ vượt ngưỡng — chỉ trả về nhóm có từ 2 item trở lên
    (nhóm 1 item không phải "đi cùng nhau", bỏ qua khi diễn giải bằng lời)."""
    loadings = result.loadings
    dominant_factor = loadings.abs().idxmax(axis=1)
    groups: dict[str, list[str]] = {col: [] for col in loadings.columns}
    for row_code in rows:
        factor_col = dominant_factor.loc[row_code]
        if abs(loadings.loc[row_code, factor_col]) >= threshold:
            groups[factor_col].append(row_code)
    return [items for items in groups.values() if len(items) >= 2]


def _add_factor_section(doc: Document, result: FactorAnalysisResult, rows: list[str], tmp: Path) -> None:
    doc.add_heading(result.label, level=3)
    if result.low_ratio_warning:
        doc.add_paragraph(
            "Số việc/vấn đề được hỏi trong mục này khá nhiều so với số phiếu (85) — kết quả "
            "dưới đây chỉ mang tính minh hoạ xu hướng, chưa đủ chắc chắn để kết luận dứt khoát."
        ).runs[0].italic = True

    groupings = _factor_groupings(result, rows)
    if groupings:
        doc.add_paragraph(
            f"Trong số {result.n_items} việc/vấn đề được hỏi, một số việc có xu hướng đi cùng "
            f"nhau trong cùng 1 phiếu (phiếu nào ghi vợ làm/quyết việc này thì cũng thường ghi "
            f"vợ làm/quyết những việc kia):"
        )
        for items in groupings:
            labels = [_item_label(result.qid, code) for code in items]
            doc.add_paragraph(", ".join(labels) + ".", style="List Bullet")
    else:
        doc.add_paragraph("Không tìm thấy nhóm việc/vấn đề nào rõ rệt đi cùng nhau trong mục này.")

    png = charts.scree_plot(result.eigenvalues, tmp / f"scree_{result.qid}.png", f"Scree plot — {result.label}")
    _add_picture(doc, png, width_cm=13)


def _cluster_highlights(cluster: ClusterResult, cluster_id: int, top_n: int = 4) -> list[str]:
    """So đặc điểm 1 cụm với mức trung bình chung toàn mẫu, lấy top_n đặc trưng lệch
    nhiều nhất — diễn giải bằng "cao/thấp hơn mức chung", KHÔNG in điểm số thô (đúng §9:
    không viết kiểu "cụm 2 có điểm tự chủ trung bình 0.62")."""
    overall = cluster.feature_matrix_raw.mean()
    means = cluster.cluster_means_raw.loc[cluster_id]
    diffs = (means - overall).abs().sort_values(ascending=False)
    highlights = []
    for feat in diffs.index[:top_n]:
        direction = "cao hơn" if means[feat] > overall[feat] else "thấp hơn"
        highlights.append(f"{feat.rstrip(' (%)')} {direction} mức chung")
    return highlights


def _add_advanced_section(doc: Document, advanced_info: dict, tmp: Path) -> None:
    doc.add_heading("Tầng 7 — Phân tích nâng cao", level=1)
    doc.add_paragraph(
        "Phần này dùng 2 kỹ thuật phân tích sâu hơn để tìm cấu trúc ẩn trong dữ liệu: việc "
        "nào thường đi cùng việc nào (phân tích nhân tố), và 85 phụ nữ khảo sát có chia "
        "được thành vài nhóm có đặc điểm khác nhau hay không (phân cụm). Với cỡ mẫu 85 "
        "phiếu, cả 2 kỹ thuật này cho kết quả mang tính GỢI Ý/THAM KHẢO, không phải phân "
        "loại chính thức hay kết luận chắc chắn."
    )

    doc.add_heading("Các việc/vấn đề thường đi cùng nhau", level=2)
    factor_info = advanced_info["factor"]
    _add_factor_section(doc, factor_info["Q14"], Q14_ROWS, tmp)
    _add_factor_section(doc, factor_info["Q32"], Q32_ROWS, tmp)

    doc.add_heading("Phân nhóm phụ nữ khảo sát theo mức độ tự chủ", level=2)
    cluster = advanced_info["cluster"]
    total = sum(cluster.cluster_sizes.values())
    doc.add_paragraph(
        f"Dựa trên nhiều chỉ số cùng lúc (sở hữu thiết bị, đi lại tự chủ, tham gia đoàn "
        f"thể/tập huấn, vay vốn, đứng tên đất, mức tham gia phân công lao động và ra quyết "
        f"định), {total} phiếu được chia thành {cluster.best_k} nhóm có đặc điểm khác nhau "
        f"rõ nhất."
    )
    for cid in sorted(cluster.cluster_sizes):
        n = cluster.cluster_sizes[cid]
        pct = n / total * 100
        doc.add_heading(f"Nhóm {cid + 1} ({n} phiếu, {pct:.0f}%)", level=3)
        highlights = _cluster_highlights(cluster, cid)
        doc.add_paragraph("Đặc điểm nổi bật so với mức chung: " + "; ".join(highlights) + ".")

    png = charts.cluster_scatter(cluster.pca_2d, cluster.labels, tmp / "cluster_scatter_docx.png", "Phân cụm 85 phiếu")
    _add_picture(doc, png, width_cm=13)
    doc.add_paragraph(
        "Lưu ý: cỡ mẫu 85 phiếu khá nhỏ để phân cụm ổn định — kết quả này chỉ mang tính "
        "gợi ý, có thể dùng để tham khảo khi thiết kế hoạt động can thiệp theo nhóm đối "
        "tượng, không nên dùng để xếp loại chính thức từng người."
    ).runs[0].italic = True
