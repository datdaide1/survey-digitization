"""Báo cáo DOCX — bản 4 TRỤ CỘT (docs/implement-plan-statistics-and-client-report.md, 26/07),
viết trên góc nhìn chuyên gia nghiên cứu thị trường/chuỗi giá trị dược liệu của dự án, không
phải người đọc SPSS thô. Văn phong theo §3 (giữ nguyên từ bản gốc §9): ngôn ngữ đời thường
gắn thực tế trồng trọt/gia đình, không thuật ngữ thống kê trong narrative, không quy kết/đánh
giá thiếu sót, "liên hệ/quan sát được" thay vì quan hệ nhân quả, luôn nhắc n=85 là mẫu nhỏ.
"""

from __future__ import annotations

import tempfile
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from . import mpl_charts as charts
from . import pillars

TITLE_COLOR = RGBColor(0x1F, 0x4E, 0x78)


def _heading(doc: Document, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = TITLE_COLOR
    return h


def _add_page_numbers(doc: Document) -> None:
    """Footer 'Trang X/Y' — chi tiết chuyên nghiệp nhỏ hay bị bỏ quên ở báo cáo tự sinh
    (26/07 tối, phản hồi khách "làm đẹp/chuyên nghiệp hơn"). Dùng field code PAGE/NUMPAGES
    thô (python-docx không có API cấp cao cho việc này) — Word tính lại khi mở; LibreOffice
    hiện đúng ngay khi build vì đã render qua recalc/convert lúc kiểm tra."""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.clear()

    def _field(field_code: str):
        run = p.add_run()
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = f" {field_code} "
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_begin)
        run._r.append(instr)
        run._r.append(fld_end)

    prefix = p.add_run("Trang ")
    prefix.font.size = Pt(9)
    prefix.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    _field("PAGE")
    mid = p.add_run("/")
    mid.font.size = Pt(9)
    mid.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    _field("NUMPAGES")


def _toc_entry(doc: Document, label: str) -> None:
    p = doc.add_paragraph(f"• {label}")
    p.paragraph_format.space_after = Pt(2)


def _note(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def _picture(doc: Document, png_path: str, width_cm: float = 15) -> None:
    doc.add_picture(png_path, width=Cm(width_cm))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def _simple_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        table.rows[0].cells[j].text = h
    for r in rows:
        cells = table.add_row().cells
        for j, v in enumerate(r):
            cells[j].text = str(v)


def _freq_rows(freq: dict[str, Any]) -> list[list[str]]:
    return [[r["label"], str(r["n"]), f"{r['pct']:.0f}%"] for r in freq["rows"] if r["n"] > 0]


def build_pillar_docx(df: pd.DataFrame, out_path: str | Path, tmp_dir: str | None = None) -> None:
    data = pillars.compute_all(df)
    tmp = Path(tmp_dir) if tmp_dir else Path(tempfile.mkdtemp(prefix="pillar_charts_"))
    n_total = len(df)

    doc = Document()
    for style_name in ("Normal",):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(11)

    title = doc.add_heading("Thị trường, chuỗi giá trị, rào cản sản xuất và môi trường chính sách đối với cây dược liệu", level=0)
    for run in title.runs:
        run.font.color.rgb = TITLE_COLOR
    doc.add_paragraph(
        "Báo cáo dựa trên 85 phiếu khảo sát bán cấu trúc với phụ nữ trồng/kinh doanh cây dược "
        "liệu tại Lào Cai và Lai Châu — một phần của dự án tìm hiểu thị trường, chuỗi giá trị, "
        "các rào cản sản xuất và môi trường chính sách đối với cây dược liệu tại Việt Nam."
    )
    _note(
        doc,
        "Mẫu khảo sát 85 phiếu, thu thập ở cấp hộ/phụ nữ — kết quả mang tính gợi ý cho vùng khảo "
        "sát, không đại diện thống kê cho toàn bộ vùng trồng dược liệu. Mọi con số trong báo cáo "
        "là liên hệ quan sát được từ dữ liệu, không phải kết luận nhân quả.",
    )
    _add_page_numbers(doc)

    doc.add_paragraph("Nội dung báo cáo").runs[0].bold = True
    for label in (
        "A. Thị trường & mức độ gắn bó với cây dược liệu",
        "B. Vị trí trong chuỗi giá trị dược liệu (gồm học thức, lãnh đạo, đi lại, thiết bị số)",
        "C. Rào cản sản xuất (gồm rào cản theo dân tộc)",
        "D. Môi trường chính sách/thể chế (gồm SWOT, kênh vay vốn theo dân tộc)",
        "E. Chỉ số vai trò trong chuỗi giá trị dược liệu (gồm theo nhóm tuổi kết hôn)",
    ):
        _toc_entry(doc, label)

    # ------------------------------------------------------------------
    # A. Thị trường & mức độ gắn bó
    # ------------------------------------------------------------------
    _heading(doc, "A. Thị trường & mức độ gắn bó với cây dược liệu")
    a = data["A"]
    top_income_src = max(a["q7"]["rows"], key=lambda r: r["n"])
    doc.add_paragraph(
        f"{a['q7']['rows'][0]['pct']:.0f}% phụ nữ được khảo sát coi cây dược liệu là một nguồn "
        "thu nhập chính của gia đình. Bên cạnh đó, phần lớn hộ vẫn duy trì song song các nguồn "
        f"thu khác (trồng trọt {next((r['pct'] for r in a['q7']['rows'] if r['code']=='trong_trot'), 0):.0f}%, "
        "phi nông nghiệp, chăn nuôi...) — cây dược liệu thường là một phần, chưa phải toàn bộ "
        "sinh kế của hộ."
    )
    img = charts.multi_select_bar("Nguồn thu nhập chính (Q7)", [r["label"] for r in a["q7"]["rows"]], [r["pct"] for r in a["q7"]["rows"]], n_total, tmp / "a_q7.png")
    _picture(doc, img)

    doc.add_paragraph("Tỷ lệ thu nhập từ dược liệu trong tổng thu nhập gia đình:")
    _simple_table(doc, ["Tỷ lệ thu nhập từ dược liệu", "Số phiếu", "Tỷ lệ %"], _freq_rows(a["q8"]))

    desc = a["q9_desc"]
    if desc["n"]:
        doc.add_paragraph(
            f"Số năm gắn bó với cây dược liệu trung bình là {desc['mean']:.0f} năm (trung vị "
            f"{desc['median']:.0f} năm), người mới bắt đầu ít nhất {desc['min']:.0f} năm, người lâu "
            f"nhất {desc['max']:.0f} năm — cho thấy đây không phải một hoạt động mới với đa số hộ."
        )

    doc.add_paragraph("So sánh giữa nhóm mới trồng (dưới 1 năm) và nhóm đã gắn bó lâu hơn (từ 1 năm trở lên):")
    for group in a["experience_x_income"]:
        rows = [[c["label"], str(c["n"]), f"{c['pct']:.0f}%"] for c in group["cells"]]
        doc.add_paragraph(f"{group['group_label']} (n={group['n_group']})").runs[0].bold = True
        _simple_table(doc, ["Tỷ lệ thu nhập từ dược liệu", "Số phiếu", "Tỷ lệ % trong nhóm"], rows)
    _note(doc, "Mẫu số mỗi bảng nhỏ là cỡ nhóm kinh nghiệm đó, không phải 85 — so sánh giữa 2 nhóm, không phải % trên toàn mẫu.")

    # ------------------------------------------------------------------
    # B. Vị trí chuỗi giá trị
    # ------------------------------------------------------------------
    _heading(doc, "B. Vị trí trong chuỗi giá trị dược liệu")
    b = data["B"]
    top_node = max(b["q30"]["rows"], key=lambda r: r["n"])
    doc.add_paragraph(
        f"Phụ nữ được khảo sát tham gia nhiều nhất ở khâu \"{top_node['label'].lower()}\" "
        f"({top_node['pct']:.0f}%), là khâu đầu chuỗi, giá trị gia tăng thường thấp hơn so với "
        "khâu chế biến/thương mại/tiêu thụ. Nhiều phiếu tham gia cùng lúc nhiều khâu (ví dụ vừa "
        "sản xuất vừa thu hái), nên tỷ lệ các khâu có thể cộng lại vượt 100%."
    )
    img = charts.bar_chart(b["q30"], tmp / "b_q30.png")
    _picture(doc, img)

    doc.add_paragraph(
        "Trong số phiếu có tham gia từng khâu, tỷ lệ thu nhập từ dược liệu phân bố như sau "
        "(khâu thương mại/tiêu thụ so với khâu sản xuất/thu hái):"
    )
    key_nodes = {"san_xuat", "thuong_mai", "tieu_thu"}
    for node in b["node_x_income"]:
        if node["node_code"] not in key_nodes or node["n_group"] == 0:
            continue
        rows = [[c["label"], str(c["n"]), f"{c['pct']:.0f}%"] for c in node["cells"]]
        doc.add_paragraph(f"Khâu {node['node_label'].lower()} (n={node['n_group']})").runs[0].bold = True
        _simple_table(doc, ["Tỷ lệ thu nhập từ dược liệu", "Số phiếu", "Tỷ lệ % trong nhóm"], rows)

    doc.add_paragraph(
        "So \"ai làm\" (Q14 — người trực tiếp thực hiện công việc) với \"ai quyết\" (Q32 — người "
        "có tiếng nói quyết định) cho cùng một chủ đề:"
    )
    pair_rows = [[p["title"], f"{p['lam_pct']:.0f}%", f"{p['quyet_pct']:.0f}%"] for p in b["lam_vs_quyet"]]
    _simple_table(doc, ["Chủ đề", "Vợ/cả hai LÀM (%)", "Vợ/cùng QUYẾT ĐỊNH (%)"], pair_rows)
    img = charts.grouped_bar(
        "Ai làm so với ai quyết (cùng chủ đề)",
        [p["title"][:40] for p in b["lam_vs_quyet"]], ["Làm", "Quyết định"],
        [{"Làm": p["lam_pct"], "Quyết định": p["quyet_pct"]} for p in b["lam_vs_quyet"]],
        tmp / "b_lam_quyet.png",
    )
    _picture(doc, img)
    doc.add_paragraph(
        "Ở cả 3 chủ đề, tỷ lệ \"có quyết định\" không thấp hơn tỷ lệ \"trực tiếp làm\" — cho thấy "
        "phụ nữ trong mẫu khảo sát này không chỉ là người thực hiện công việc mà cũng có tiếng "
        "nói trong quyết định liên quan, ít nhất theo tự báo cáo của người trả lời."
    )

    # ------------------------------------------------------------------
    # Bổ sung 26/07 (tối) — học thức/lãnh đạo/đi lại/thiết bị (docs implement plan §8)
    # ------------------------------------------------------------------
    _heading(doc, "Học thức, vai trò lãnh đạo, đi lại độc lập và thiết bị số", level=2)
    doc.add_paragraph(
        "Bốn đặc điểm dưới đây không phải rào cản/chính sách hay chuỗi giá trị nói chung, "
        "nhưng liên hệ trực tiếp tới việc phụ nữ có tiến được lên khâu \"cao giá\" (thương "
        "mại/tiêu thụ) hay không — nên đặt chung ở đây thay vì chỉ nằm rời rạc ở tần suất "
        "trải phẳng."
    )

    def _cao_gia_pct(group: dict[str, Any]) -> float | None:
        if not group["q30"]:
            return None
        by_code = {r["code"]: r["pct"] for r in group["q30"]["rows"]}
        return by_code.get("thuong_mai", 0.0) + by_code.get("tieu_thu", 0.0)

    edu_rows = [
        [g["group_label"], str(g["n_group"]), f"{_cao_gia_pct(g):.0f}%" if _cao_gia_pct(g) is not None else "n/a"]
        for g in b["education_x_q30"]
    ]
    doc.add_paragraph("Học thức (Q5) — tỷ lệ tham gia khâu thương mại và/hoặc tiêu thụ theo nhóm học vấn:")
    _simple_table(doc, ["Nhóm học vấn", "Số phiếu", "Tham gia khâu thương mại/tiêu thụ"], edu_rows)

    lead_rows = [
        [g["group_label"], str(g["n_group"]), f"{_cao_gia_pct(g):.0f}%" if _cao_gia_pct(g) is not None else "n/a"]
        for g in b["leadership_x_q30"]
    ]
    doc.add_paragraph("Vai trò lãnh đạo nhóm sản xuất/HTX/quản lý rừng (Q33):")
    _simple_table(doc, ["Nhóm", "Số phiếu", "Tham gia khâu thương mại/tiêu thụ"], lead_rows)
    _note(doc, "Nhóm 'có vai trò lãnh đạo' rất nhỏ (n≈5/85) — chênh lệch quan sát được, không phải bằng chứng thống kê chắc chắn.")

    mob_rows = [
        [g["group_label"], str(g["n_group"]), f"{_cao_gia_pct(g):.0f}%" if _cao_gia_pct(g) is not None else "n/a"]
        for g in b["mobility_x_q30"]
    ]
    doc.add_paragraph("Có xe máy riêng để tự đi lại (Q25):")
    _simple_table(doc, ["Nhóm", "Số phiếu", "Tham gia khâu thương mại/tiêu thụ"], mob_rows)

    dev = b["device_ownership_vs_usage"]
    doc.add_paragraph(
        f"{dev['own_pct']:.0f}% phụ nữ được khảo sát sở hữu điện thoại thông minh — tỷ lệ sở "
        "hữu cao, nhưng tỷ lệ dùng thiết bị đó cho mục đích kinh tế thấp hơn nhiều:"
    )
    dev_rows = [["Sở hữu điện thoại thông minh (vợ)", f"{dev['own_pct']:.0f}%"]] + [
        [r["label"], f"{r['pct']:.0f}%"] for r in dev["rows"]
    ]
    _simple_table(doc, ["", "Tỷ lệ % (trên 85 phiếu)"], dev_rows)
    doc.add_paragraph(
        "Khoảng cách giữa \"có thiết bị\" và \"dùng thiết bị để bán hàng/quảng bá\" gợi ý dư "
        "địa hỗ trợ kỹ năng số/thương mại điện tử, không chỉ dừng ở phát triển hạ tầng/thiết bị."
    )

    # ------------------------------------------------------------------
    # C. Rào cản sản xuất
    # ------------------------------------------------------------------
    _heading(doc, "C. Rào cản sản xuất")
    c = data["C"]
    top_barrier = c["q28"]["rows"][0]
    doc.add_paragraph(
        f"Rào cản được nhắc đến nhiều nhất là \"{top_barrier['label'].lower()}\" "
        f"({top_barrier['pct']:.0f}%). Chi tiết từng rào cản:"
    )
    img = charts.multi_select_bar("Khó khăn khi trồng/kinh doanh dược liệu (Q28)", [r["label"] for r in c["q28"]["rows"]], [r["pct"] for r in c["q28"]["rows"]], n_total, tmp / "c_q28.png")
    _picture(doc, img)

    doc.add_paragraph(
        "So sánh tỷ lệ từng rào cản giữa nhóm chỉ tham gia khâu sản xuất/thu hái/chế biến và "
        "nhóm có tham gia thêm khâu thương mại/tiêu thụ:"
    )
    node_defs = c["node_group_defs"]
    header = ["Rào cản"] + [label for _c, label in node_defs]
    rows = []
    for barrier in c["barrier_x_node_group"]:
        row = [barrier["label"]]
        for cell in barrier["cells"]:
            row.append(f"{cell['pct']:.0f}% (n={cell['n']}/{cell['n_group']})")
        rows.append(row)
    _simple_table(doc, header, rows)
    doc.add_paragraph(
        "Đây là liên hệ quan sát được trên mẫu 85 phiếu, không phải bằng chứng nhân quả — nhưng "
        "nếu một rào cản xuất hiện rõ rệt nhiều hơn ở nhóm 'chỉ khâu sản xuất/thu hái/chế biến', "
        "đó là gợi ý đáng cân nhắc cho việc thiết kế hỗ trợ giúp phụ nữ tiến lên khâu có giá trị "
        "cao hơn trong chuỗi."
    )

    eth_barrier = c["barrier_x_ethnicity"]
    top_barrier_row = eth_barrier[0]
    doc.add_paragraph(
        "Rào cản theo dân tộc (giữ nguyên từng dân tộc quan sát được — Kinh, Dao, Mông, Nùng, "
        f"Tày — theo đúng yêu cầu khách, không gộp thành 'dân tộc thiểu số'): với '{top_barrier_row['label'].lower()}', "
        + ", ".join(f"{cell['group']} {cell['pct']:.0f}% (n={cell['n']}/{cell['n_group']})" for cell in top_barrier_row["cells"])
        + ". Chi tiết đủ 7 rào cản × 5 dân tộc xem sheet \"C. Rào cản\" trong file Excel."
    )
    _note(doc, "Kinh (n=4), Tày (n=4), Nùng (n=2) là nhóm rất nhỏ trong mẫu này — % chỉ mang tính tham khảo, không đại diện thống kê cho cả dân tộc đó.")

    # ------------------------------------------------------------------
    # D. Môi trường chính sách
    # ------------------------------------------------------------------
    _heading(doc, "D. Môi trường chính sách/thể chế")
    d = data["D"]
    doc.add_paragraph("Tiếp cận vốn vay sản xuất/kinh doanh trong 5 năm qua:")
    _simple_table(doc, ["Nguồn vốn", "Số phiếu", "Tỷ lệ %"], _freq_rows(d["q22a"]))
    doc.add_paragraph("Tiếp cận tập huấn và hỗ trợ khác:")
    _simple_table(
        doc, ["Nội dung", "Số phiếu", "Tỷ lệ %"],
        _freq_rows(d["q21a"]) + _freq_rows(d["q23"]) + _freq_rows(d["q11"]),
    )

    eth_support = d["support_x_ethnicity"]
    chua_row = next((r for r in eth_support if r["code"] == "chua"), None)
    if chua_row:
        doc.add_paragraph(
            "Tỷ lệ 'chưa từng vay vốn' theo dân tộc (giữ nguyên từng dân tộc, không gộp): "
            + ", ".join(f"{cell['group']} {cell['pct']:.0f}% (n={cell['n']}/{cell['n_group']})" for cell in chua_row["cells"])
            + ". Chi tiết đủ các kênh vay vốn × 5 dân tộc xem sheet \"D. Chính sách\" trong file Excel."
        )
        _note(doc, "Kinh (n=4), Tày (n=4), Nùng (n=2) là nhóm rất nhỏ trong mẫu này — % chỉ mang tính tham khảo.")

    _heading(doc, "Tổng hợp SWOT — môi trường chính sách", level=2)
    swot_table = doc.add_table(rows=2, cols=2)
    swot_table.style = "Light Grid Accent 1"
    swot_labels = [("ĐIỂM MẠNH", "strengths"), ("CƠ HỘI", "opportunities"), ("ĐIỂM YẾU", "weaknesses"), ("THÁCH THỨC", "threats")]
    for i, (title_txt, key) in enumerate(swot_labels):
        cell = swot_table.rows[i // 2].cells[i % 2]
        cell.paragraphs[0].add_run(title_txt).bold = True
        for bullet in d["swot"][key]:
            cell.add_paragraph(f"• {bullet}")

    # ------------------------------------------------------------------
    # E. Chỉ số vai trò trong chuỗi giá trị
    # ------------------------------------------------------------------
    _heading(doc, "E. Chỉ số vai trò trong chuỗi giá trị dược liệu")
    e = data["E"]
    for qid, r in (("Q14", e["Q14"]), ("Q32", e["Q32"])):
        doc.add_paragraph(r.label).runs[0].bold = True
        if r.composite is not None:
            mean_v = float(r.composite.mean())
            doc.add_paragraph(
                f"Trung bình toàn mẫu: {mean_v:.0f}/100 — điểm càng cao nghĩa là phụ nữ càng có "
                "vai trò rõ trong các việc/quyết định thuộc nhóm này."
            )
            img = charts.composite_distribution(f"Phân bố chỉ số — {r.label}", r.composite.tolist(), tmp / f"e_{qid}.png")
            _picture(doc, img)

            marriage_breakdown = e.get(f"{qid}_x_marriage_age")
            if marriage_breakdown:
                rows = [[g["group_label"], str(g["n_group"]), f"{g['mean']:.0f}/100" if g["mean"] is not None else "n/a"] for g in marriage_breakdown]
                doc.add_paragraph("Theo nhóm tuổi kết hôn (Q6):")
                _simple_table(doc, ["Nhóm tuổi kết hôn", "Số phiếu", "Chỉ số trung bình"], rows)
                vals = {g["group_code"]: g["mean"] for g in marriage_breakdown if g["mean"] is not None}
                if "<18" in vals and ">=18" in vals:
                    higher = "kết hôn trước 18 tuổi (tảo hôn)" if vals["<18"] > vals[">=18"] else "kết hôn từ 18 tuổi"
                    # 26/07 (đêm, phản hồi khách — "chỗ Q6 bị duplicate"): câu văn cũ không
                    # nhắc số cụ thể nên khi Q14 và Q32 cùng chiều (nhóm tảo hôn cao hơn ở cả
                    # 2), 2 đoạn văn ra Y HỆT NHAU, đọc như copy nhầm — giờ chèn thẳng số của
                    # ĐÚNG chỉ số đang nói (r.label) để 2 đoạn luôn khác nhau.
                    doc.add_paragraph(
                        f"Ở {r.label.lower()}, nhóm {higher} có chỉ số trung bình cao hơn "
                        f"({vals['<18']:.0f}/100 so với {vals['>=18']:.0f}/100) — quan sát này "
                        "KHÔNG theo chiều giả thuyết \"kết hôn sớm đi cùng tiếng nói thấp hơn\" "
                        "thường gặp trong nghiên cứu trao quyền phụ nữ; chỉ nêu đúng như quan "
                        "sát được trên 85 phiếu, không suy diễn nguyên nhân."
                    )
        else:
            doc.add_paragraph(
                "Các dòng câu hỏi trong nhóm này chưa đủ nhất quán để gộp thành 1 chỉ số duy nhất "
                "— giữ nguyên kết quả từng dòng riêng (xem sheet tương ứng trong file Excel)."
            )

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
