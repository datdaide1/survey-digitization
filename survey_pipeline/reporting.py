"""Generic XLSX and DOCX reports for arbitrary flattened survey data."""

from __future__ import annotations

from pathlib import Path
import json
from typing import Any

from .config import ProjectConfig


def _frequency_tables(frame):
    tables = []
    for column in frame.columns:
        if column == "record_id":
            continue
        counts = frame[column].fillna("(missing)").astype(str).value_counts(dropna=False)
        tables.append((column, counts))
    return tables


def build_reports(config: ProjectConfig):
    import pandas as pd
    from docx import Document
    from openpyxl import Workbook
    from openpyxl.styles import Font

    frame = pd.read_csv(config.paths.combined)
    config.paths.reports.mkdir(parents=True, exist_ok=True)
    stem = str(config.reporting.get("filename") or config.project_id)
    xlsx_path = config.paths.reports / f"{stem}.xlsx"
    docx_path = config.paths.reports / f"{stem}.docx"

    wb = Workbook()
    data = wb.active
    data.title = "Data"
    data.append(list(frame.columns))
    for cell in data[1]:
        cell.font = Font(bold=True)
    for row in frame.itertuples(index=False, name=None):
        data.append(list(row))
    freq = wb.create_sheet("Frequencies")
    freq.append(["field", "value", "count", "percent"])
    for column, counts in _frequency_tables(frame):
        denominator = int(counts.sum()) or 1
        for value, count in counts.items():
            freq.append([column, value, int(count), float(count) / denominator])

    analysis_path = config.paths.stats / "analysis.json"
    advanced = json.loads(analysis_path.read_text(encoding="utf-8")) if analysis_path.is_file() else None
    if advanced:
        descriptive = wb.create_sheet("Descriptives")
        descriptive.append(["variable", "n", "mean", "std", "min", "q1", "median", "q3", "max"])
        for variable, values in advanced["descriptives"].items():
            descriptive.append([variable, *(values.get(key) for key in ("n", "mean", "std", "min", "q1", "median", "q3", "max"))])

        associations = wb.create_sheet("Associations")
        associations.append(["variable_a", "variable_b", "method", "effect", "p_value", "adjusted_p_value", "n", "min_group_n", "caution"])
        for item in advanced["associations"]:
            associations.append([item.get(key) for key in ("variable_a", "variable_b", "method", "effect", "p_value", "adjusted_p_value", "n", "min_group_n", "caution")])

        reliability = wb.create_sheet("Reliability")
        reliability.append(["scale", "alpha", "n", "items", "reason"])
        for scale, values in advanced["reliability"].items():
            reliability.append([scale, values.get("alpha"), values.get("n"), values.get("items"), values.get("reason")])

        group_tests = wb.create_sheet("Group Tests")
        group_tests.append(["outcome", "group", "method", "statistic", "p_value", "caution", "group_details"])
        for item in advanced["group_comparisons"]:
            group_tests.append([item.get("outcome"), item.get("group"), item.get("method"), item.get("statistic"), item.get("p_value"), item.get("caution"), json.dumps(item.get("groups"), ensure_ascii=False)])

        methods = wb.create_sheet("Methodology")
        methods.append(["setting", "value"])
        for key, value in advanced["policy"].items():
            methods.append([key, str(value)])
    wb.save(xlsx_path)

    document = Document()
    document.add_heading(config.title, 0)
    document.add_paragraph(f"Records: {len(frame)}")
    for column, counts in _frequency_tables(frame):
        document.add_heading(column, level=1)
        table = document.add_table(rows=1, cols=3)
        table.rows[0].cells[0].text = "Value"
        table.rows[0].cells[1].text = "Count"
        table.rows[0].cells[2].text = "Percent"
        denominator = int(counts.sum()) or 1
        for value, count in counts.items():
            cells = table.add_row().cells
            cells[0].text = str(value)
            cells[1].text = str(int(count))
            cells[2].text = f"{float(count) / denominator:.1%}"
    if advanced:
        document.add_heading("Statistical analysis", level=1)
        document.add_paragraph(
            "Configured analyses use variable-appropriate effect measures, report sample sizes, "
            "flag small groups, and adjust simultaneous association tests when configured. "
            "Associations are observational and must not be interpreted as causal effects."
        )
        document.add_paragraph(
            f"Association pairs: {len(advanced['associations'])}; "
            f"reliability scales: {len(advanced['reliability'])}; "
            f"group comparisons: {len(advanced['group_comparisons'])}."
        )
    document.save(docx_path)
    return xlsx_path, docx_path
